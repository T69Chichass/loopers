"""
Document processing pipeline for extracting text from various file formats
and preparing them for vector storage and querying.
"""
import os
import uuid
import hashlib
import logging
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
from datetime import datetime
import asyncio

# PDF processing
import fitz  # PyMuPDF
from pdfplumber import PDF

# Document processing
from docx import Document as DocxDocument
import pandas as pd

# Text processing
import re
from sentence_transformers import SentenceTransformer
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.corpus import stopwords

# Database and vector store
from sqlalchemy.orm import Session
from dependencies import get_database_manager, get_pinecone_manager, get_embedding_manager
from database import DocumentMetadata, DocumentChunk
from exceptions import *

# Download required NLTK data
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt')

try:
    nltk.data.find('corpora/stopwords')
except LookupError:
    nltk.download('stopwords')

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Main document processing class for handling various file formats."""
    
    def __init__(self):
        self.db_manager = get_database_manager()
        self.pinecone_manager = get_pinecone_manager()
        self.embedding_manager = get_embedding_manager()
        self.supported_formats = {'.pdf', '.docx', '.doc', '.txt', '.csv'}
        self.chunk_size = 1000  # characters per chunk
        self.chunk_overlap = 200  # overlap between chunks
    
    async def process_document(
        self, 
        file_path: str, 
        document_type: str = "policy",
        category: str = "general",
        metadata: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """
        Process a document end-to-end: extract text, chunk, embed, and store.
        
        Args:
            file_path: Path to the document file
            document_type: Type of document (insurance, legal, hr, etc.)
            category: Document category
            metadata: Additional metadata
        
        Returns:
            Processing result summary
        """
        start_time = datetime.utcnow()
        document_id = str(uuid.uuid4())
        
        try:
            logger.info(f"Starting document processing: {file_path}")
            
            # Validate file
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"Document not found: {file_path}")
            
            file_path = Path(file_path)
            if file_path.suffix.lower() not in self.supported_formats:
                raise ValueError(f"Unsupported file format: {file_path.suffix}")
            
            # Create document metadata record
            doc_metadata = await self._create_document_metadata(
                document_id, file_path, document_type, category, metadata
            )
            
            # Extract text from document
            logger.info(f"Extracting text from {file_path.name}")
            extracted_text = await self._extract_text(file_path)
            
            if not extracted_text.strip():
                raise ValueError("No text content extracted from document")
            
            # Create text chunks
            logger.info(f"Creating text chunks for document {document_id}")
            chunks = await self._create_chunks(extracted_text, document_id)
            
            # Generate embeddings and store in Pinecone
            logger.info(f"Generating embeddings for {len(chunks)} chunks")
            embedding_results = await self._process_embeddings(chunks, document_id)
            
            # Update document metadata
            await self._update_document_status(
                document_id, "completed", len(chunks), datetime.utcnow()
            )
            
            processing_time = (datetime.utcnow() - start_time).total_seconds()
            
            result = {
                "document_id": document_id,
                "filename": file_path.name,
                "document_type": document_type,
                "category": category,
                "text_length": len(extracted_text),
                "chunk_count": len(chunks),
                "embeddings_created": len(embedding_results),
                "processing_time": processing_time,
                "status": "completed",
                "created_at": start_time.isoformat()
            }
            
            logger.info(f"Document processing completed: {document_id}")
            return result
            
        except Exception as e:
            logger.error(f"Document processing failed: {e}")
            await self._update_document_status(document_id, "failed", 0, None, str(e))
            raise DocumentProcessingError(f"Failed to process document: {e}")
    
    async def _extract_text(self, file_path: Path) -> str:
        """Extract text from various file formats."""
        extension = file_path.suffix.lower()
        
        try:
            if extension == '.pdf':
                return await self._extract_text_from_pdf(file_path)
            elif extension in ['.docx', '.doc']:
                return await self._extract_text_from_docx(file_path)
            elif extension == '.txt':
                return await self._extract_text_from_txt(file_path)
            elif extension == '.csv':
                return await self._extract_text_from_csv(file_path)
            else:
                raise ValueError(f"Unsupported file format: {extension}")
                
        except Exception as e:
            logger.error(f"Text extraction failed for {file_path}: {e}")
            raise
    
    async def _extract_text_from_pdf(self, file_path: Path) -> str:
        """Extract text from PDF files using multiple methods for best results."""
        text_content = []
        
        try:
            # Method 1: PyMuPDF (faster, good for most PDFs)
            doc = fitz.open(file_path)
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text = page.get_text()
                if text.strip():
                    text_content.append(f"[Page {page_num + 1}]\n{text}")
            doc.close()
            
            # If no text extracted, try pdfplumber (better for complex layouts)
            if not any(content.strip() for content in text_content):
                logger.info(f"Trying alternative PDF extraction for {file_path.name}")
                text_content = []
                
                with PDF.open(file_path) as pdf:
                    for page_num, page in enumerate(pdf.pages):
                        text = page.extract_text()
                        if text and text.strip():
                            text_content.append(f"[Page {page_num + 1}]\n{text}")
            
            return "\n\n".join(text_content)
            
        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            raise
    
    async def _extract_text_from_docx(self, file_path: Path) -> str:
        """Extract text from DOCX files."""
        try:
            doc = DocxDocument(file_path)
            text_content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            return "\n\n".join(text_content)
            
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            raise
    
    async def _extract_text_from_txt(self, file_path: Path) -> str:
        """Extract text from TXT files."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try different encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read()
                except UnicodeDecodeError:
                    continue
            raise ValueError("Could not decode text file")
    
    async def _extract_text_from_csv(self, file_path: Path) -> str:
        """Extract text from CSV files."""
        try:
            df = pd.read_csv(file_path)
            text_content = []
            
            # Add column headers
            text_content.append("Columns: " + " | ".join(df.columns))
            
            # Add each row as text
            for index, row in df.iterrows():
                row_text = []
                for col, value in row.items():
                    if pd.notna(value):
                        row_text.append(f"{col}: {value}")
                if row_text:
                    text_content.append(" | ".join(row_text))
            
            return "\n\n".join(text_content)
            
        except Exception as e:
            logger.error(f"CSV extraction failed: {e}")
            raise
    
    async def _create_chunks(self, text: str, document_id: str) -> List[Dict[str, Any]]:
        """Create overlapping text chunks for better context preservation."""
        # Clean the text
        text = self._clean_text(text)
        
        # Split into sentences for better chunk boundaries
        sentences = sent_tokenize(text)
        
        chunks = []
        current_chunk = ""
        current_chunk_sentences = []
        chunk_index = 0
        
        for sentence in sentences:
            # Check if adding this sentence would exceed chunk size
            test_chunk = current_chunk + " " + sentence if current_chunk else sentence
            
            if len(test_chunk) <= self.chunk_size:
                current_chunk = test_chunk
                current_chunk_sentences.append(sentence)
            else:
                # Save current chunk if it has content
                if current_chunk.strip():
                    chunk_id = f"{document_id}_chunk_{chunk_index}"
                    chunks.append({
                        "chunk_id": chunk_id,
                        "document_id": document_id,
                        "chunk_text": current_chunk.strip(),
                        "chunk_index": chunk_index,
                        "chunk_size": len(current_chunk),
                        "sentence_count": len(current_chunk_sentences)
                    })
                    chunk_index += 1
                
                # Start new chunk with overlap
                overlap_sentences = current_chunk_sentences[-2:] if len(current_chunk_sentences) > 2 else current_chunk_sentences
                current_chunk = " ".join(overlap_sentences + [sentence])
                current_chunk_sentences = overlap_sentences + [sentence]
        
        # Add final chunk
        if current_chunk.strip():
            chunk_id = f"{document_id}_chunk_{chunk_index}"
            chunks.append({
                "chunk_id": chunk_id,
                "document_id": document_id,
                "chunk_text": current_chunk.strip(),
                "chunk_index": chunk_index,
                "chunk_size": len(current_chunk),
                "sentence_count": len(current_chunk_sentences)
            })
        
        return chunks
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text content."""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep punctuation
        text = re.sub(r'[^\w\s.,!?;:()\-\'"']', ' ', text)
        
        # Remove extra spaces
        text = re.sub(r'\s+', ' ', text)
        
        return text.strip()
    
    async def _process_embeddings(self, chunks: List[Dict[str, Any]], document_id: str) -> List[Dict[str, Any]]:
        """Generate embeddings and store in Pinecone."""
        results = []
        
        # Store chunks in database first
        db = self.db_manager.get_db()
        try:
            for chunk_data in chunks:
                # Create database record
                db_chunk = DocumentChunk(
                    chunk_id=chunk_data["chunk_id"],
                    document_id=document_id,
                    chunk_text=chunk_data["chunk_text"],
                    chunk_index=chunk_data["chunk_index"],
                    chunk_size=chunk_data["chunk_size"]
                )
                db.add(db_chunk)
            db.commit()
            
            # Generate embeddings
            texts = [chunk["chunk_text"] for chunk in chunks]
            embeddings = self.embedding_manager.encode_batch(texts)
            
            # Prepare vectors for Pinecone
            vectors_to_upsert = []
            for i, (chunk_data, embedding) in enumerate(zip(chunks, embeddings)):
                vector_data = {
                    "id": chunk_data["chunk_id"],
                    "values": embedding,
                    "metadata": {
                        "text": chunk_data["chunk_text"],
                        "document_id": document_id,
                        "chunk_index": chunk_data["chunk_index"],
                        "chunk_size": chunk_data["chunk_size"]
                    }
                }
                vectors_to_upsert.append(vector_data)
                
                # Update database record
                db_chunk = db.query(DocumentChunk).filter(
                    DocumentChunk.chunk_id == chunk_data["chunk_id"]
                ).first()
                if db_chunk:
                    db_chunk.embedding_generated = True
                    db_chunk.pinecone_id = chunk_data["chunk_id"]
            
            # Batch upsert to Pinecone
            if vectors_to_upsert:
                self.pinecone_manager.index.upsert(vectors=vectors_to_upsert)
                logger.info(f"Upserted {len(vectors_to_upsert)} vectors to Pinecone")
            
            db.commit()
            results = [{"chunk_id": chunk["chunk_id"], "status": "success"} for chunk in chunks]
            
        except Exception as e:
            db.rollback()
            logger.error(f"Error processing embeddings: {e}")
            raise
        finally:
            self.db_manager.close_db(db)
        
        return results
    
    async def _create_document_metadata(
        self, 
        document_id: str, 
        file_path: Path, 
        document_type: str,
        category: str, 
        metadata: Optional[Dict[str, Any]]
    ) -> DocumentMetadata:
        """Create document metadata record in database."""
        db = self.db_manager.get_db()
        try:
            # Get file stats
            file_stats = os.stat(file_path)
            
            doc_metadata = DocumentMetadata(
                document_id=document_id,
                title=metadata.get("title", file_path.stem) if metadata else file_path.stem,
                document_type=document_type,
                category=category,
                filename=file_path.name,
                file_path=str(file_path),
                file_size=file_stats.st_size,
                processing_status="processing"
            )
            
            db.add(doc_metadata)
            db.commit()
            return doc_metadata
            
        except Exception as e:
            db.rollback()
            logger.error(f"Error creating document metadata: {e}")
            raise
        finally:
            self.db_manager.close_db(db)
    
    async def _update_document_status(
        self, 
        document_id: str, 
        status: str, 
        chunk_count: int,
        processed_at: Optional[datetime],
        error_message: Optional[str] = None
    ):
        """Update document processing status."""
        db = self.db_manager.get_db()
        try:
            doc = db.query(DocumentMetadata).filter(
                DocumentMetadata.document_id == document_id
            ).first()
            
            if doc:
                doc.processing_status = status
                doc.chunk_count = chunk_count
                if processed_at:
                    doc.processed_at = processed_at
                db.commit()
                
        except Exception as e:
            db.rollback()
            logger.error(f"Error updating document status: {e}")
        finally:
            self.db_manager.close_db(db)
    
    async def delete_document(self, document_id: str) -> Dict[str, Any]:
        """Delete document and all associated chunks from database and Pinecone."""
        db = self.db_manager.get_db()
        try:
            # Get document chunks
            chunks = db.query(DocumentChunk).filter(
                DocumentChunk.document_id == document_id
            ).all()
            
            # Delete from Pinecone
            chunk_ids = [chunk.chunk_id for chunk in chunks]
            if chunk_ids:
                self.pinecone_manager.index.delete(ids=chunk_ids)
                logger.info(f"Deleted {len(chunk_ids)} vectors from Pinecone")
            
            # Delete from database
            db.query(DocumentChunk).filter(
                DocumentChunk.document_id == document_id
            ).delete()
            
            doc = db.query(DocumentMetadata).filter(
                DocumentMetadata.document_id == document_id
            ).first()
            
            if doc:
                db.delete(doc)
                db.commit()
                return {
                    "document_id": document_id,
                    "status": "deleted",
                    "chunks_deleted": len(chunk_ids)
                }
            else:
                return {"error": "Document not found"}
                
        except Exception as e:
            db.rollback()
            logger.error(f"Error deleting document: {e}")
            raise
        finally:
            self.db_manager.close_db(db)
    
    async def get_document_status(self, document_id: str) -> Optional[Dict[str, Any]]:
        """Get document processing status and metadata."""
        db = self.db_manager.get_db()
        try:
            doc = db.query(DocumentMetadata).filter(
                DocumentMetadata.document_id == document_id
            ).first()
            
            if doc:
                chunk_count = db.query(DocumentChunk).filter(
                    DocumentChunk.document_id == document_id
                ).count()
                
                return {
                    "document_id": doc.document_id,
                    "title": doc.title,
                    "document_type": doc.document_type,
                    "category": doc.category,
                    "filename": doc.filename,
                    "file_size": doc.file_size,
                    "processing_status": doc.processing_status,
                    "chunk_count": chunk_count,
                    "created_at": doc.created_at.isoformat(),
                    "processed_at": doc.processed_at.isoformat() if doc.processed_at else None
                }
            return None
            
        except Exception as e:
            logger.error(f"Error getting document status: {e}")
            return None
        finally:
            self.db_manager.close_db(db)


class DocumentProcessingError(Exception):
    """Exception raised when document processing fails."""
    pass
